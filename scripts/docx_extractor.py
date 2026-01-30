"""
DOCX Text Extraction Module

Extracts text from DOCX files using python-docx.
"""

from pathlib import Path
from typing import List, Optional

import docx
from scripts.utils import logger


class DocxExtractor:
    """Extract and process text from DOCX files."""

    def __init__(self, docx_path: Path):
        self.docx_path = Path(docx_path)
        if not self.docx_path.exists():
            raise FileNotFoundError(f"DOCX not found: {docx_path}")

        try:
            self.doc = docx.Document(self.docx_path)
        except Exception as e:
            raise RuntimeError(f"Failed to open DOCX file: {e}")

    def extract_all(self, skip_pages: Optional[List[int]] = None) -> str:
        """
        Extract all text from the DOCX.

        Args:
            skip_pages: Ignored for DOCX as it doesn't have fixed pages like PDF.

        Returns:
            Extracted text as a single string
        """
        if skip_pages:
            logger.warning("skip_pages argument is ignored for DOCX files")

        text_parts = []
        
        # Extract text from paragraphs
        logger.info(f"Extracting text from {len(self.doc.paragraphs)} paragraphs...")
        
        for paragraph in self.doc.paragraphs:
            text = paragraph.text.strip()
            if text:
                text_parts.append(text)
                
        return "\n\n".join(text_parts)

    def close(self) -> None:
        """No-op for python-docx but kept for API consistency."""
        pass

    def __enter__(self) -> "DocxExtractor":
        return self

    def __exit__(self, exc_type, exc_val, exc_tb) -> None:
        self.close()
