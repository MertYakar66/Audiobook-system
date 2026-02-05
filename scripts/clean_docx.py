"""
DOCX Chapter Cleaner for The Intelligent Investor

Extracts only the main chapters from the DOCX file,
removing non-essential front matter, footnotes, and commentary annotations.
"""

from docx import Document
from docx.shared import Pt
import re
from pathlib import Path

def analyze_document(input_path: str):
    """Analyze the document structure and find chapters."""
    doc = Document(input_path)
    paragraphs = [(i, p.text.strip()) for i, p in enumerate(doc.paragraphs) if p.text.strip()]
    
    print(f"Total paragraphs with content: {len(paragraphs)}")
    print("\n=== Looking for chapter headings ===\n")
    
    chapters = []
    for idx, (para_idx, text) in enumerate(paragraphs):
        # Match chapter headings
        if re.match(r'^CHAPTER\s+\d+', text, re.IGNORECASE):
            chapters.append((para_idx, text[:80]))
            print(f"  Found: {text[:80]}")
    
    return chapters

def extract_clean_chapters(input_path: str, output_path: str):
    """
    Extract only main chapters (1-20), removing:
    - Title pages and front matter
    - Jason Zweig's commentary sections
    - Footnotes and references
    - Index and appendices that aren't readable
    - Duplicate chapters from end of book
    """
    doc = Document(input_path)
    paragraphs = doc.paragraphs
    
    # Patterns to identify sections to skip
    skip_patterns = [
        r'^\*\s',  # Footnote markers
        r'^†\s',   # Footnote markers
        r'^\d+\s+Coauthored',  # Footnote text
        r'^Commentary on Chapter',  # Zweig's commentary headers
        r'^COMMENTARY ON CHAPTER',
        r'^\s*\d+\s*$',  # Page numbers only
        r'^About the Authors',
        r'^About the Publisher',
        r'^Credits',
        r'^Front Cover',
        r'^Copyright',
        r'^Contents$',
        r'^Epigraph',
    ]
    
    # Find where main content starts (after front matter)
    content_start = None
    for i, p in enumerate(paragraphs):
        text = p.text.strip()
        if re.match(r'^CHAPTER\s+1\b', text, re.IGNORECASE):
            content_start = i
            print(f"Found Chapter 1 at paragraph {i}")
            break
    
    if content_start is None:
        print("ERROR: Could not find Chapter 1!")
        return
    
    # Create new document with only main chapters
    new_doc = Document()
    
    # Add title
    title = new_doc.add_heading('The Intelligent Investor', 0)
    new_doc.add_paragraph('by Benjamin Graham')
    new_doc.add_paragraph('')
    
    current_chapter = 0
    skip_until_next_chapter = False
    paragraphs_added = 0
    chapters_found = []
    finished = False
    
    for i in range(content_start, len(paragraphs)):
        if finished:
            break
            
        text = paragraphs[i].text.strip()
        
        if not text:
            continue
            
        # Check if this is a chapter heading
        chapter_match = re.match(r'^CHAPTER\s+(\d+)', text, re.IGNORECASE)
        if chapter_match:
            chapter_num = int(chapter_match.group(1))
            
            # Stop if we've seen this chapter before (duplicates at end)
            if chapter_num <= current_chapter and current_chapter >= 20:
                print(f"Stopping at duplicate Chapter {chapter_num} (already processed up to {current_chapter})")
                finished = True
                break
            
            # Also stop if going backwards significantly
            if chapter_num < current_chapter and current_chapter > 0:
                print(f"Stopping at Chapter {chapter_num} (going backwards from {current_chapter})")
                finished = True
                break
                
            current_chapter = chapter_num
            chapters_found.append(current_chapter)
            skip_until_next_chapter = False
            
            # Add chapter heading
            new_doc.add_heading(f'Chapter {current_chapter}', 1)
            paragraphs_added += 1
            continue
        
        # Check if this is Zweig's commentary (skip it)
        if re.match(r'^Commentary on Chapter', text, re.IGNORECASE):
            skip_until_next_chapter = True
            continue
        
        if skip_until_next_chapter:
            continue
            
        # Check if we should skip this paragraph
        should_skip = False
        for pattern in skip_patterns:
            if re.match(pattern, text, re.IGNORECASE):
                should_skip = True
                break
        
        if should_skip:
            continue
        
        # Check for end of main content (appendices, index, etc.)
        if re.match(r'^(APPENDIX|INDEX|ACKNOWLEDGMENTS)', text, re.IGNORECASE):
            break
        
        # Add the paragraph
        new_doc.add_paragraph(text)
        paragraphs_added += 1
    
    # Save the cleaned document
    new_doc.save(output_path)
    
    print(f"\n=== Extraction Complete ===")
    print(f"Chapters found: {chapters_found}")
    print(f"Paragraphs added: {paragraphs_added}")
    print(f"Output saved to: {output_path}")
    
    return chapters_found

if __name__ == "__main__":
    input_file = r"input/The Intelligent Investor.docx"
    output_file = r"input/The_Intelligent_Investor_Clean.docx"
    
    print("=== Analyzing Document ===\n")
    analyze_document(input_file)
    
    print("\n=== Extracting Clean Chapters ===\n")
    extract_clean_chapters(input_file, output_file)
